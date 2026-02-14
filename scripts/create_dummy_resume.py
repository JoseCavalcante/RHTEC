from docx import Document

doc = Document()
doc.add_heading('Currículo - Maria Exemplo', 0)

doc.add_paragraph('Nome: Maria Exemplo')
doc.add_paragraph('Senioridade: Pleno')
doc.add_paragraph('Experiência: 4 anos')
doc.add_paragraph('Skills: Python, FastAPI, React, SQL')
doc.add_paragraph('Resumo: Desenvolvedora Full Stack com foco em entregas rápidas e código limpo.')

doc.save('dummy_resume.docx')
print("dummy_resume.docx created")
